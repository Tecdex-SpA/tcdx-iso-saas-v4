from fastapi import FastAPI, Header, HTTPException, Response
from pydantic import BaseModel, Field
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError
from tempfile import NamedTemporaryFile
import base64
import csv
import ipaddress
import io
import json
import os
import socket

from app.routes.ai import router as ai_router
from app.routes.audit_documents import router as audit_documents_router
from app.routes.senior_auditor_v2 import router as senior_auditor_v2_router
from app.core.config import settings
from app.core.db import test_db_connection

docs_options = {}
if not getattr(settings, "AI_ENGINE_PUBLIC_DOCS", False):
    docs_options = {
        "docs_url": None,
        "redoc_url": None,
        "openapi_url": None,
    }

app = FastAPI(title=settings.APP_NAME, **docs_options)

MAX_REMOTE_FILE_BYTES = int(os.getenv("AI_REMOTE_FILE_MAX_BYTES", str(15 * 1024 * 1024)))
ALLOWED_REMOTE_SCHEMES = {"http", "https"}


# =========================================================
# Helpers generales
# =========================================================

def get_configured_ai_token() -> Optional[str]:
    return (
        getattr(settings, "AI_INTERNAL_TOKEN", None)
        or os.getenv("AI_INTERNAL_TOKEN")
        or os.getenv("OWN_AI_SHARED_SECRET")
        or None
    )


def require_ai_token(
    x_ai_token: Optional[str] = None,
    x_internal_token: Optional[str] = None,
) -> None:
    configured = get_configured_ai_token()
    if not configured:
        raise HTTPException(status_code=503, detail="AI token not configured")

    provided = x_internal_token or x_ai_token
    if provided != configured:
        raise HTTPException(status_code=401, detail="Invalid AI token")


@app.get("/")
def root():
    """
    Endpoint raíz liviano para validar que el AI Engine está activo.
    No expone estado interno, secretos ni conexión a base de datos.
    """
    return {
        "ok": True,
        "service": settings.APP_NAME,
        "status": "running",
    }


@app.head("/")
def root_head():
    """
    Permite validar con curl -I http://host:8001 sin recibir 404.
    """
    return Response(status_code=200)


@app.get("/health")
def health():
    """
    Health check público básico para balanceadores, monitoreo simple o curl.
    """
    return {
        "ok": True,
        "service": settings.APP_NAME,
        "status": "running",
    }


@app.head("/health")
def health_head():
    return Response(status_code=200)


@app.get("/health/deep")
def health_deep(
    x_ai_token: Optional[str] = Header(default=None),
    x_internal_token: Optional[str] = Header(default=None),
):
    """
    Health check profundo protegido por token interno.
    Valida token y conexión a PostgreSQL sin exponer secretos.
    """
    require_ai_token(
        x_ai_token=x_ai_token,
        x_internal_token=x_internal_token,
    )

    return {
        "ok": True,
        "service": settings.APP_NAME,
        "status": "running",
        "db_ok": test_db_connection(),
        "ai_token_configured": bool(get_configured_ai_token()),
    }


def safe_dict(value: Any) -> Dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}
    return {}


def safe_list(value: Any) -> List[Any]:
    return value if isinstance(value, list) else []


def to_text(value: Any) -> str:
    return str(value or "").strip()


def normalize_for_match(value: Any) -> str:
    return to_text(value).lower()


def contains_any(text: str, terms: List[str]) -> bool:
    haystack = normalize_for_match(text)
    return any(term in haystack for term in terms)


def clamp_score(value: float) -> float:
    return max(0, min(100, round(value, 2)))


def infer_file_type(file_name: Optional[str], mime_type: Optional[str]) -> str:
    name = to_text(file_name).lower()
    mime = to_text(mime_type).lower()

    if name.endswith(".pdf") or mime == "application/pdf":
        return "pdf"
    if name.endswith(".docx") or "wordprocessingml" in mime:
        return "docx"
    if name.endswith(".xlsx") or "spreadsheetml" in mime:
        return "xlsx"
    if name.endswith(".csv") or mime == "text/csv":
        return "csv"
    if name.endswith(".txt") or mime.startswith("text/plain"):
        return "txt"
    if name.endswith(".json") or mime == "application/json":
        return "json"
    if name.endswith(".png") or name.endswith(".jpg") or name.endswith(".jpeg") or name.endswith(".webp") or mime.startswith("image/"):
        return "image"
    return "binary"


def validate_remote_file_url(file_url: str) -> Optional[str]:
    parsed = urlparse(file_url)

    if parsed.scheme.lower() not in ALLOWED_REMOTE_SCHEMES:
        return "Esquema de URL no permitido"

    host = parsed.hostname
    if not host:
        return "Host de URL inválido"

    try:
        addresses = socket.getaddrinfo(host, None)
    except socket.gaierror:
        return "No fue posible resolver el host"

    for address in addresses:
        ip_text = address[4][0]
        try:
            ip = ipaddress.ip_address(ip_text)
        except ValueError:
            return "Dirección IP inválida"

        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_multicast
            or ip.is_reserved
            or ip.is_unspecified
        ):
            return "Host de URL no permitido"

    return None


def fetch_remote_file(file_url: Optional[str], timeout_seconds: int = 15) -> Dict[str, Any]:
    if not file_url:
        return {"ok": False, "error": "No file_url provided", "content": b"", "content_type": None, "source": "none"}

    validation_error = validate_remote_file_url(file_url)
    if validation_error:
        return {
            "ok": False,
            "error": validation_error,
            "content": b"",
            "content_type": None,
            "source": "remote_download",
        }

    try:
        request = Request(
            file_url,
            headers={"User-Agent": "Tecdex-AI-Engine/1.0"},
        )
        with urlopen(request, timeout=timeout_seconds) as response:
            content = response.read(MAX_REMOTE_FILE_BYTES + 1)
            if len(content) > MAX_REMOTE_FILE_BYTES:
                return {
                    "ok": False,
                    "error": "Archivo remoto excede el tamaño máximo permitido",
                    "content": b"",
                    "content_type": response.headers.get("Content-Type"),
                    "status": getattr(response, "status", 200),
                    "source": "remote_download",
                }

            return {
                "ok": True,
                "content": content,
                "content_type": response.headers.get("Content-Type"),
                "status": getattr(response, "status", 200),
                "source": "remote_download",
            }
    except HTTPError as exc:
        return {
            "ok": False,
            "error": f"HTTPError {exc.code}",
            "content": b"",
            "content_type": None,
            "source": "remote_download",
        }
    except URLError as exc:
        return {
            "ok": False,
            "error": f"URLError {exc.reason}",
            "content": b"",
            "content_type": None,
            "source": "remote_download",
        }
    except Exception as exc:
        return {
            "ok": False,
            "error": str(exc),
            "content": b"",
            "content_type": None,
            "source": "remote_download",
        }


def decode_text_bytes(content: bytes) -> str:
    if not content:
        return ""

    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except Exception:
            continue

    return ""


def get_inline_file_bytes(evidence: Dict[str, Any]) -> Dict[str, Any]:
    base64_content = to_text(evidence.get("file_content_base64"))
    encoding = to_text(evidence.get("file_content_encoding")).lower()

    if not base64_content:
        return {
            "ok": False,
            "content": b"",
            "content_type": evidence.get("file_mime_type"),
            "error": "No inline file provided",
            "source": "inline_payload",
        }

    if encoding not in ("base64", ""):
        return {
            "ok": False,
            "content": b"",
            "content_type": evidence.get("file_mime_type"),
            "error": f"Unsupported inline encoding: {encoding}",
            "source": "inline_payload",
        }

    try:
        content = base64.b64decode(base64_content)
        return {
            "ok": True,
            "content": content,
            "content_type": evidence.get("file_mime_type"),
            "error": None,
            "source": "inline_payload",
        }
    except Exception as exc:
        return {
            "ok": False,
            "content": b"",
            "content_type": evidence.get("file_mime_type"),
            "error": f"Inline decode failed: {exc}",
            "source": "inline_payload",
        }


def extract_text_from_pdf(content: bytes) -> Dict[str, Any]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return {
            "ok": False,
            "text": "",
            "parser": None,
            "page_count": None,
            "warning": "pypdf no instalado",
        }

    try:
        with NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()

            reader = PdfReader(tmp.name)
            pages = []
            for page in reader.pages:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    pages.append("")

            text = "\n".join([p for p in pages if p]).strip()
            return {
                "ok": True,
                "text": text,
                "parser": "pypdf",
                "page_count": len(reader.pages),
                "warning": None,
            }
    except Exception as exc:
        return {
            "ok": False,
            "text": "",
            "parser": "pypdf",
            "page_count": None,
            "warning": str(exc),
        }


def extract_text_from_docx(content: bytes) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return {
            "ok": False,
            "text": "",
            "parser": None,
            "warning": "python-docx no instalado",
        }

    try:
        with NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()

            doc = Document(tmp.name)
            parts: List[str] = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            return {
                "ok": True,
                "text": "\n".join(parts).strip(),
                "parser": "python-docx",
                "warning": None,
            }
    except Exception as exc:
        return {
            "ok": False,
            "text": "",
            "parser": "python-docx",
            "warning": str(exc),
        }


def extract_text_from_xlsx(content: bytes) -> Dict[str, Any]:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception:
        return {
            "ok": False,
            "text": "",
            "parser": None,
            "sheet_count": None,
            "warning": "openpyxl no instalado",
        }

    try:
        with NamedTemporaryFile(suffix=".xlsx", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()

            wb = load_workbook(tmp.name, data_only=True)
            lines: List[str] = []

            for ws in wb.worksheets:
                lines.append(f"[HOJA] {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    values = [str(v).strip() for v in row if v is not None and str(v).strip()]
                    if values:
                        lines.append(" | ".join(values))

            return {
                "ok": True,
                "text": "\n".join(lines).strip(),
                "parser": "openpyxl",
                "sheet_count": len(wb.worksheets),
                "warning": None,
            }
    except Exception as exc:
        return {
            "ok": False,
            "text": "",
            "parser": "openpyxl",
            "sheet_count": None,
            "warning": str(exc),
        }


def extract_text_from_csv(content: bytes) -> Dict[str, Any]:
    try:
        decoded = decode_text_bytes(content)
        if not decoded:
            return {
                "ok": False,
                "text": "",
                "parser": "csv",
                "warning": "No se pudo decodificar CSV",
            }

        sio = io.StringIO(decoded)
        reader = csv.reader(sio)
        lines = []
        for row in reader:
            values = [str(v).strip() for v in row if str(v).strip()]
            if values:
                lines.append(" | ".join(values))

        return {
            "ok": True,
            "text": "\n".join(lines).strip(),
            "parser": "csv",
            "warning": None,
        }
    except Exception as exc:
        return {
            "ok": False,
            "text": "",
            "parser": "csv",
            "warning": str(exc),
        }


def extract_text_from_json(content: bytes) -> Dict[str, Any]:
    decoded = decode_text_bytes(content)
    if not decoded:
        return {"ok": False, "text": "", "parser": "json", "warning": "No se pudo decodificar JSON"}

    try:
        obj = json.loads(decoded)
        pretty = json.dumps(obj, ensure_ascii=False, indent=2)
        return {"ok": True, "text": pretty, "parser": "json", "warning": None}
    except Exception as exc:
        return {"ok": False, "text": decoded, "parser": "json", "warning": str(exc)}


def extract_text_from_image(content: bytes) -> Dict[str, Any]:
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except Exception:
        return {
            "ok": False,
            "text": "",
            "parser": None,
            "warning": "OCR no disponible (Pillow/pytesseract no instalados)",
        }

    try:
        with NamedTemporaryFile(suffix=".png", delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            image = Image.open(tmp.name)
            text = pytesseract.image_to_string(image, lang="spa+eng").strip()
            return {
                "ok": True,
                "text": text,
                "parser": "pytesseract",
                "warning": None,
            }
    except Exception as exc:
        return {
            "ok": False,
            "text": "",
            "parser": "pytesseract",
            "warning": str(exc),
        }


def best_effort_extraction(payload: Dict[str, Any]) -> Dict[str, Any]:
    evidence = safe_dict(payload.get("evidence"))
    current_extract = safe_dict(payload.get("current_extract"))

    file_name = evidence.get("file_name")
    mime_type = evidence.get("file_mime_type")
    file_url = evidence.get("file_url")
    file_type = infer_file_type(file_name, mime_type)

    existing_raw_text = to_text(current_extract.get("raw_text"))
    if existing_raw_text:
        return {
            "extraction_status": "completed",
            "extraction_engine": "reuse_current_extract",
            "file_type": file_type,
            "mime_type": mime_type,
            "raw_text": existing_raw_text,
            "structured_json": {
                "source": "current_extract",
                "reused": True,
            },
            "ocr_used": bool(current_extract.get("ocr_used")),
            "detected_language": current_extract.get("detected_language") or "es",
            "page_count": current_extract.get("page_count"),
            "sheet_count": current_extract.get("sheet_count"),
            "image_count": current_extract.get("image_count"),
            "extraction_notes": None,
        }

    inline_file = get_inline_file_bytes(evidence)
    if inline_file.get("ok"):
        fetched = inline_file
    else:
        fetched = fetch_remote_file(file_url)

    content = fetched.get("content", b"") if fetched.get("ok") else b""
    effective_mime = mime_type or fetched.get("content_type")

    extraction_notes = []
    parser_used = None
    page_count = None
    sheet_count = None
    image_count = None
    ocr_used = False
    raw_text = ""

    if file_type in ("txt", "binary") and content:
        raw_text = decode_text_bytes(content)
        parser_used = "plain_decode"

    elif file_type == "json" and content:
        parsed = extract_text_from_json(content)
        raw_text = parsed.get("text", "")
        parser_used = parsed.get("parser")
        if parsed.get("warning"):
            extraction_notes.append(parsed["warning"])

    elif file_type == "csv" and content:
        parsed = extract_text_from_csv(content)
        raw_text = parsed.get("text", "")
        parser_used = parsed.get("parser")
        if parsed.get("warning"):
            extraction_notes.append(parsed["warning"])

    elif file_type == "pdf" and content:
        parsed = extract_text_from_pdf(content)
        raw_text = parsed.get("text", "")
        parser_used = parsed.get("parser")
        page_count = parsed.get("page_count")
        if parsed.get("warning"):
            extraction_notes.append(parsed["warning"])

    elif file_type == "docx" and content:
        parsed = extract_text_from_docx(content)
        raw_text = parsed.get("text", "")
        parser_used = parsed.get("parser")
        if parsed.get("warning"):
            extraction_notes.append(parsed["warning"])

    elif file_type == "xlsx" and content:
        parsed = extract_text_from_xlsx(content)
        raw_text = parsed.get("text", "")
        parser_used = parsed.get("parser")
        sheet_count = parsed.get("sheet_count")
        if parsed.get("warning"):
            extraction_notes.append(parsed["warning"])

    elif file_type == "image" and content:
        image_count = 1
        parsed = extract_text_from_image(content)
        raw_text = parsed.get("text", "")
        parser_used = parsed.get("parser")
        ocr_used = bool(parsed.get("parser"))
        if parsed.get("warning"):
            extraction_notes.append(parsed["warning"])

    if not raw_text:
        description = to_text(evidence.get("description"))
        control = safe_dict(payload.get("control"))
        operation = safe_dict(payload.get("operation"))

        fallback_parts = [
            description,
            to_text(control.get("description")),
            to_text(control.get("clause")),
            to_text(control.get("standard_code")),
            to_text(operation.get("operation_name")),
        ]
        raw_text = "\n".join([p for p in fallback_parts if p]).strip()

    extraction_status = "completed" if raw_text else "limited"
    if file_type == "image" and ocr_used and raw_text:
        extraction_status = "completed_with_ocr"

    if not fetched.get("ok"):
        if fetched.get("source") == "inline_payload":
            extraction_notes.append(f"No se pudo decodificar archivo inline: {fetched.get('error')}")
        elif file_url:
            extraction_notes.append(f"No se pudo descargar archivo: {fetched.get('error')}")

    if not parser_used and content:
        extraction_notes.append("No hubo parser especializado disponible para este tipo de archivo.")

    if not raw_text:
        extraction_notes.append("No fue posible extraer texto utilizable; se usó solo contexto mínimo.")

    return {
        "extraction_status": extraction_status,
        "extraction_engine": parser_used or "heuristic_extractor",
        "file_type": file_type,
        "mime_type": effective_mime,
        "raw_text": raw_text,
        "structured_json": {
            "source": fetched.get("source"),
            "download_ok": fetched.get("ok", False),
            "download_error": fetched.get("error"),
            "byte_size": len(content) if content else 0,
            "parser_used": parser_used,
            "inline_used": fetched.get("source") == "inline_payload" and fetched.get("ok", False),
        },
        "ocr_used": ocr_used,
        "detected_language": "es",
        "page_count": page_count,
        "sheet_count": sheet_count,
        "image_count": image_count,
        "extraction_notes": " | ".join([n for n in extraction_notes if n]) or None,
    }


def build_entities(payload: Dict[str, Any], extraction: Dict[str, Any]) -> List[str]:
    evidence = safe_dict(payload.get("evidence"))
    control = safe_dict(payload.get("control"))
    operation = safe_dict(payload.get("operation"))
    action_plan = safe_dict(payload.get("action_plan"))

    entities = [
        to_text(evidence.get("file_name")),
        to_text(control.get("standard_code")),
        to_text(control.get("clause")),
        to_text(control.get("description")),
        to_text(operation.get("operation_name")),
        to_text(action_plan.get("title")),
    ]

    if extraction.get("file_type"):
        entities.append(f"tipo:{extraction.get('file_type')}")

    return [e for e in entities if e]


def build_assessment(payload: Dict[str, Any], extraction: Dict[str, Any]) -> Dict[str, Any]:
    evidence = safe_dict(payload.get("evidence"))
    control = safe_dict(payload.get("control"))
    operation = safe_dict(payload.get("operation"))
    lifecycle = safe_dict(payload.get("lifecycle"))

    status = to_text(evidence.get("status")).lower()
    file_name = to_text(evidence.get("file_name"))
    standard_code = to_text(control.get("standard_code"))
    clause = to_text(control.get("clause"))
    control_description = to_text(control.get("description"))
    operation_name = to_text(operation.get("operation_name"))
    extracted_text = to_text(extraction.get("raw_text"))
    combined_text = "\n".join([
        extracted_text,
        to_text(evidence.get("description")),
        control_description,
        to_text(action_plan.get("title")) if (action_plan := safe_dict(payload.get("action_plan"))) else "",
    ])

    has_real_file = bool(
        evidence.get("file_name")
        and (
            evidence.get("file_url")
            or evidence.get("file_content_base64")
        )
    )
    has_text = len(extracted_text) > 40

    evidence_coverage_pct = float(lifecycle.get("evidence_coverage_pct") or 0)
    avg_health_score = float(lifecycle.get("avg_health_score") or 0)
    health_status = to_text(control.get("tenant_control_health_status") or lifecycle.get("health_status")).lower()

    has_date_signal = contains_any(combined_text, [
        "fecha", "vigencia", "vigente", "actualizado", "aprobado", "aprobada",
        "revisado", "revisada", "2024", "2025", "2026"
    ])
    has_owner_signal = contains_any(combined_text, [
        "responsable", "owner", "aprobador", "auditor", "gerente", "jefe",
        "encargado", "ciso", "coordinador"
    ])
    has_result_signal = contains_any(combined_text, [
        "resultado", "cumplimiento", "evidencia", "registro", "acta",
        "verificación", "verificacion", "eficacia", "seguimiento", "prueba",
        "simulacro", "revisión", "revision"
    ])
    if contains_any(combined_text, [
        "sin fecha", "ni fecha", "no tiene fecha", "sin vigencia", "no indica vigencia",
        "sin revision", "sin revisión", "no indica revision", "no indica revisión"
    ]):
        has_date_signal = False

    if contains_any(combined_text, [
        "sin responsable", "ni responsable", "no tiene responsable", "sin aprobador", "ni aprobador",
        "no indica responsable", "no indica aprobador", "sin dueño", "sin dueno"
    ]):
        has_owner_signal = False

    if contains_any(combined_text, [
        "sin resultado", "ni resultado", "no tiene resultado", "sin evidencia de ejecucion",
        "sin evidencia de ejecución", "no demuestra ejecucion", "no demuestra ejecución",
        "sin verificacion", "sin verificación"
    ]):
        has_result_signal = False
    has_control_signal = bool(control_description and control_description.lower()[:28] in combined_text.lower()) or contains_any(
        combined_text,
        [term for term in [standard_code.lower(), clause.lower(), operation_name.lower()] if term]
    )

    pertinence_score = 55
    sufficiency_score = 35
    freshness_score = 65
    traceability_score = 45
    consistency_score = 55
    compliance_impact_score = 40

    if standard_code:
        pertinence_score += 15
    if clause:
        pertinence_score += 10
    if control_description:
        pertinence_score += 5

    if has_real_file:
        traceability_score += 25

    if has_text:
        sufficiency_score += 30
        consistency_score += 10
        compliance_impact_score += 20
    else:
        compliance_impact_score += 5

    if status == "aprobada":
        sufficiency_score += 15
        traceability_score += 10
        compliance_impact_score += 20
    elif status == "rechazada":
        sufficiency_score -= 10
        compliance_impact_score -= 20

    if evidence_coverage_pct < 30:
        freshness_score -= 5

    if avg_health_score >= 80:
        compliance_impact_score += 5

    if has_date_signal:
        freshness_score += 12
        traceability_score += 6
    else:
        freshness_score -= 12
        sufficiency_score -= 6

    if has_owner_signal:
        traceability_score += 10
        consistency_score += 5
    else:
        traceability_score -= 8

    if has_result_signal:
        sufficiency_score += 10
        compliance_impact_score += 8
    else:
        sufficiency_score -= 8

    if has_control_signal:
        pertinence_score += 8
        consistency_score += 8
    else:
        pertinence_score -= 10
        consistency_score -= 5

    if health_status in ("deteriorado", "critical", "red", "rojo"):
        compliance_impact_score += 8
    elif health_status in ("saludable", "healthy", "green", "verde"):
        compliance_impact_score += 3

    pertinence_score = clamp_score(pertinence_score)
    sufficiency_score = clamp_score(sufficiency_score)
    freshness_score = clamp_score(freshness_score)
    traceability_score = clamp_score(traceability_score)
    consistency_score = clamp_score(consistency_score)
    compliance_impact_score = clamp_score(compliance_impact_score)

    strong_evidence = (
        has_real_file
        and has_text
        and has_date_signal
        and has_owner_signal
        and has_result_signal
        and has_control_signal
        and sufficiency_score >= 75
        and traceability_score >= 70
    )

    usable_evidence = has_real_file and has_text and sufficiency_score >= 55 and traceability_score >= 55

    if status == "rechazada":
        validity_result = "no_valida"
        contribution_level = "bajo"
    elif strong_evidence and status == "aprobada":
        validity_result = "valida"
        contribution_level = "alto"
    elif usable_evidence and status in ("aprobada", "pendiente"):
        validity_result = "parcial"
        contribution_level = "medio"
    elif has_text:
        validity_result = "debil"
        contribution_level = "medio"
    else:
        validity_result = "debil"
        contribution_level = "bajo"

    risks = []
    next_steps = []

    if not has_real_file:
        risks.append("La evidencia no tiene archivo físico asociado; su trazabilidad documental es débil.")
        next_steps.append("Subir un archivo real para fortalecer la evidencia.")

    if not has_text:
        risks.append("No fue posible extraer contenido documental suficiente del archivo.")
        next_steps.append("Reprocesar con un parser/OCR especializado o adjuntar una versión legible.")

    if not has_date_signal:
        risks.append("No se identificó claramente fecha, vigencia o revisión; puede debilitarse la trazabilidad en auditoría.")
        next_steps.append("Complementar con fecha de emisión/revisión y periodo cubierto por la evidencia.")

    if not has_owner_signal:
        risks.append("No se identificó responsable o aprobador claro en la evidencia.")
        next_steps.append("Agregar responsable, aprobador o dueño del control en la evidencia o sus metadatos.")

    if not has_result_signal:
        risks.append("No se observó resultado verificable de ejecución, revisión o eficacia del control.")
        next_steps.append("Adjuntar registro de ejecución, resultado, revisión o prueba de eficacia.")

    if not has_control_signal:
        risks.append("La conexión explícita entre evidencia, norma, cláusula o control es limitada.")
        next_steps.append("Vincular la evidencia al control correcto e incluir referencia de norma/cláusula.")

    if evidence_coverage_pct < 60:
        risks.append(f"La cobertura de evidencia del estándar sigue baja ({evidence_coverage_pct:.0f}%).")
        next_steps.append("Complementar con más evidencias del mismo control u operación.")

    if status == "pendiente":
        risks.append("La evidencia aún no está aprobada; su impacto en cumplimiento todavía es provisional.")
        next_steps.append("Solicitar revisión y aprobación por auditor o administrador autorizado.")

    if status == "rechazada":
        risks.append("La evidencia fue rechazada y no debería considerarse como cierre válido del control.")
        next_steps.append("Cargar una nueva evidencia corregida y justificar su pertinencia.")

    if not risks:
        risks.append("No se observan alertas críticas inmediatas para esta evidencia.")

    if not next_steps:
        next_steps.append("Mantener vigente la evidencia y revisar periódicamente su trazabilidad y suficiencia.")

    evidence_gaps = []
    if not has_real_file:
        evidence_gaps.append("archivo_evidencia")
    if not has_text:
        evidence_gaps.append("contenido_legible")
    if not has_date_signal:
        evidence_gaps.append("fecha_vigencia_revision")
    if not has_owner_signal:
        evidence_gaps.append("responsable_aprobador")
    if not has_result_signal:
        evidence_gaps.append("resultado_verificable")
    if not has_control_signal:
        evidence_gaps.append("vinculo_control_norma")

    recommended_evidence_requests = []
    if "fecha_vigencia_revision" in evidence_gaps:
        recommended_evidence_requests.append("Solicitar versión con fecha de emisión, revisión, vigencia o periodo auditado.")
    if "responsable_aprobador" in evidence_gaps:
        recommended_evidence_requests.append("Solicitar evidencia con responsable, aprobador o dueño del control.")
    if "resultado_verificable" in evidence_gaps:
        recommended_evidence_requests.append("Solicitar registro de ejecución con resultado, muestra revisada y conclusión.")
    if "vinculo_control_norma" in evidence_gaps:
        recommended_evidence_requests.append("Solicitar referencia explícita al control, cláusula o requisito aplicable.")
    if not recommended_evidence_requests:
        recommended_evidence_requests.append("Mantener evidencia complementaria de revisión periódica y eficacia.")

    recommended_actions = []
    if validity_result in ("debil", "parcial", "no_valida"):
        recommended_actions.append("No cerrar el control automáticamente; dejarlo pendiente de revisión humana.")
        recommended_actions.append("Crear o actualizar plan de acción para cerrar las brechas documentales detectadas.")
    if health_status in ("deteriorado", "critical", "red", "rojo"):
        recommended_actions.append("Priorizar este control antes de auditoría por su salud deteriorada.")
    recommended_actions.append("Registrar decisión del auditor y criterio de aceptación o rechazo.")

    appears_complete = has_real_file and has_text
    appears_expired = False
    appears_authentic = True if has_real_file else None

    headline = (
        f"La evidencia {file_name or evidence.get('id')} "
        f"{'aporta fuertemente' if contribution_level == 'alto' else 'aporta parcialmente' if contribution_level == 'medio' else 'aporta débilmente'} "
        f"al control {clause or 'sin cláusula'}."
    )

    narrative = " ".join(
        [
            f"Norma: {standard_code or 'N/D'}.",
            f"Cláusula: {clause or 'N/D'}.",
            f"Control: {control_description or 'Sin descripción'}.",
            f"Operación: {operation_name or 'Sin operación'}.",
            f"Estado de evidencia: {status or 'N/D'}.",
            f"Archivo: {file_name or 'sin archivo físico'}.",
            f"Texto extraído: {len(extracted_text)} caracteres.",
            f"Validez estimada: {validity_result}.",
            f"Impacto esperado en cumplimiento: {compliance_impact_score:.0f}/100.",
        ]
    )

    gap_summary = (
        "La evidencia aún no cubre completamente la brecha documental del control: "
        + ", ".join(evidence_gaps)
        if validity_result in ("debil", "parcial")
        else "La evidencia cubre razonablemente el control, aunque puede fortalecerse con evidencia complementaria."
    )

    control_fit = (
        f"La evidencia se alinea con {standard_code or 'la norma'}"
        f"{' cláusula ' + clause if clause else ''}"
        f" y el control '{control_description}'."
        if control_description else None
    )

    return {
        "analysis_status": "completed_with_warnings" if extraction.get("extraction_status") == "limited" else "completed",
        "validity_result": validity_result,
        "contribution_level": contribution_level,
        "pertinence_score": pertinence_score,
        "sufficiency_score": sufficiency_score,
        "freshness_score": freshness_score,
        "traceability_score": traceability_score,
        "consistency_score": consistency_score,
        "compliance_impact_score": compliance_impact_score,
        "recommended_standard_code": standard_code or None,
        "recommended_clause": clause or None,
        "recommended_control_id": control.get("catalog_control_id") or None,
        "recommended_operation_id": operation.get("operation_id") or None,
        "headline": headline,
        "narrative": narrative,
        "risks_json": risks,
        "next_steps_json": next_steps,
        "extracted_entities_json": build_entities(payload, extraction),
        "control_fit": control_fit,
        "gap_summary": gap_summary,
        "duplicate_of_evidence_id": None,
        "appears_expired": appears_expired,
        "appears_complete": appears_complete,
        "appears_authentic": appears_authentic,
        "model_name": "own_ai_140",
        "model_version": "v1-evidence-heuristic",
        "source_system": "own_ai_140",
        "raw_response_json": {
            "heuristic": True,
            "text_length": len(extracted_text),
            "has_real_file": has_real_file,
            "status": status,
            "file_source": "inline_payload" if evidence.get("file_content_base64") else "remote_or_context",
            "auditor_reasoning": {
                "evidence_gaps": evidence_gaps,
                "signals": {
                    "has_date_signal": has_date_signal,
                    "has_owner_signal": has_owner_signal,
                    "has_result_signal": has_result_signal,
                    "has_control_signal": has_control_signal,
                    "control_health_status": health_status or None,
                },
                "recommended_evidence_requests": recommended_evidence_requests,
                "recommended_actions": recommended_actions,
                "human_approval_required": True,
            },
        },
    }


def build_chunks(extraction: Dict[str, Any], assessment: Dict[str, Any]) -> List[Dict[str, Any]]:
    text = to_text(extraction.get("raw_text"))
    if not text:
        narrative = to_text(assessment.get("narrative"))
        if not narrative:
            return []
        text = narrative

    max_size = 1200
    chunks = []
    cursor = 0
    index = 0

    while cursor < len(text):
        piece = text[cursor:cursor + max_size].strip()
        if piece:
            chunks.append(
                {
                    "chunk_index": index,
                    "chunk_type": "text",
                    "content": piece,
                    "metadata_json": {
                        "source": "evidence_process",
                        "validity_result": assessment.get("validity_result"),
                        "contribution_level": assessment.get("contribution_level"),
                    },
                }
            )
            index += 1
        cursor += max_size

    return chunks


# =========================================================
# Models
# =========================================================

class EvidenceProcessRequest(BaseModel):
    job_type: Optional[str] = Field(default="analyze_evidence")
    payload_version: Optional[int] = Field(default=1)
    request_meta: Dict[str, Any] = Field(default_factory=dict)
    evidence: Dict[str, Any] = Field(default_factory=dict)
    control: Dict[str, Any] = Field(default_factory=dict)
    operation: Dict[str, Any] = Field(default_factory=dict)
    action_plan: Dict[str, Any] = Field(default_factory=dict)
    lifecycle: Dict[str, Any] = Field(default_factory=dict)
    current_extract: Optional[Dict[str, Any]] = Field(default=None)
    current_assessment: Optional[Dict[str, Any]] = Field(default=None)


# =========================================================
# Routes
# =========================================================

@app.get("/health")
def health():
    return {
        "ok": True,
        "service": settings.APP_NAME,
        "env": settings.APP_ENV,
        "db_connection": test_db_connection(),
    }


@app.post("/api/evidences/process", tags=["AI"])
def process_evidence(
    payload: EvidenceProcessRequest,
    x_ai_token: Optional[str] = Header(default=None, alias="x-ai-token"),
    x_internal_token: Optional[str] = Header(default=None, alias="x-internal-token"),
):
    require_ai_token(x_ai_token=x_ai_token, x_internal_token=x_internal_token)

    request_payload = payload.model_dump()

    extraction = best_effort_extraction(request_payload)
    assessment = build_assessment(request_payload, extraction)
    chunks = build_chunks(extraction, assessment)

    return {
        "ok": True,
        "source": "own_ai_140",
        "job_type": payload.job_type,
        "extraction": extraction,
        "assessment": assessment,
        "chunks": chunks,
    }


app.include_router(ai_router)
app.include_router(senior_auditor_v2_router)
app.include_router(audit_documents_router)


# =========================================================
# Document AI Analysis Endpoint - TCDX Evidence Center
# =========================================================

class DocumentAnalysisRequest(BaseModel):
    tenant_id: str
    document_id: str
    file_name: Optional[str] = None
    mime_type: Optional[str] = None
    text: str = ""
    metadata: Dict[str, Any] = Field(default_factory=dict)
    active_standards: List[str] = Field(default_factory=list)
    available_controls: List[Dict[str, Any]] = Field(default_factory=list)
    instructions: Dict[str, Any] = Field(default_factory=dict)


class SemanticEvidenceAnalyzeRequest(BaseModel):
    tenant_id: str
    source_type: str
    source_id: str
    document_key: Optional[str] = None
    filename: Optional[str] = None
    title: Optional[str] = None
    text: str = ""
    metadata: Dict[str, Any] = Field(default_factory=dict)
    candidate_targets: Dict[str, List[Dict[str, Any]]] = Field(default_factory=dict)


def _doc_norm(value: Any) -> str:
    return str(value or "").strip()


def _doc_lower(value: Any) -> str:
    return _doc_norm(value).lower()


def _doc_contains_any(text: str, terms: List[str]) -> bool:
    lowered = _doc_lower(text)
    return any(term in lowered for term in terms)


def _normalize_standard_code(value: Any) -> Optional[str]:
    raw = _doc_norm(value)
    compact = raw.upper().replace(" ", "").replace("_", "").replace("-", "")

    if "27001" in compact:
        return "ISO27001:2022"
    if "9001" in compact:
        return "ISO9001:2015"
    if "42001" in compact:
        return "ISO42001:2023"
    if "22301" in compact:
        return "ISO22301:2019"
    if "14001" in compact:
        return "ISO14001:2015"
    if "37001" in compact:
        return "ISO37001:2016"
    if "50001" in compact:
        return "ISO50001:2018"

    return raw or None


def _infer_document_type(file_name: str, mime_type: str, text: str) -> str:
    haystack = f"{file_name} {mime_type} {text}".lower()

    if _doc_contains_any(haystack, ["política", "politica", "policy", "lineamientos", "directriz"]):
        return "policy"
    if _doc_contains_any(haystack, ["procedimiento", "procedure", "instrucción", "instructivo", "flujo", "paso a paso"]):
        return "procedure"
    if _doc_contains_any(haystack, ["registro", "record", "listado", "matriz", "planilla", "acta", "formulario"]):
        return "record"
    if _doc_contains_any(haystack, ["auditoría", "auditoria", "audit", "hallazgo", "no conformidad", "observación"]):
        return "audit_report"
    if _doc_contains_any(haystack, ["riesgo", "risk", "amenaza", "vulnerabilidad", "impacto", "probabilidad"]):
        return "risk_document"
    if _doc_contains_any(haystack, ["evidencia", "evidence", "respaldo", "certificado", "comprobante"]):
        return "evidence"

    return "unknown"


def _infer_standards(file_name: str, text: str, active_standards: List[str]) -> List[str]:
    haystack = f"{file_name} {text}".lower()
    detected: List[str] = []

    for candidate in active_standards or []:
        normalized = _normalize_standard_code(candidate)
        if normalized and normalized not in detected:
            raw = _doc_lower(candidate)
            compact_raw = raw.replace(":", "").replace("-", "").replace(" ", "")
            compact_haystack = haystack.replace(":", "").replace("-", "").replace(" ", "")
            if raw and (raw in haystack or compact_raw in compact_haystack):
                detected.append(normalized)

    explicit_map = [
        ("ISO9001:2015", ["iso9001", "iso 9001", "sistema de gestión de calidad", "gestion de calidad", "sgc", "calidad"]),
        ("ISO27001:2022", ["iso27001", "iso 27001", "seguridad de la información", "seguridad de la informacion", "ciberseguridad"]),
        ("ISO42001:2023", ["iso42001", "iso 42001", "inteligencia artificial", "sistema de gestión de ia", "sistema de gestion de ia"]),
        ("ISO22301:2019", ["iso22301", "iso 22301", "continuidad", "bcp", "drp", "recuperación ante desastres"]),
        ("ISO14001:2015", ["iso14001", "iso 14001", "ambiental", "medio ambiente"]),
    ]

    for standard, terms in explicit_map:
        if _doc_contains_any(haystack, terms) and standard not in detected:
            detected.append(standard)

    if not detected and active_standards:
        normalized = _normalize_standard_code(active_standards[0])
        if normalized:
            detected.append(normalized)

    return detected[:3]


def _find_probable_controls(text: str, standards: List[str], available_controls: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not available_controls:
        return []

    lowered = _doc_lower(text)
    scored = []

    for control in available_controls:
        standard = _normalize_standard_code(control.get("standard_code"))
        if standards and standard not in standards:
            continue

        clause = _doc_norm(control.get("clause"))
        control_code = _doc_norm(control.get("control_code"))
        title = _doc_norm(control.get("title"))
        description = _doc_norm(control.get("description"))

        score = 0
        reasons = []

        for field_name, field_value, weight in [
            ("clause", clause, 18),
            ("control_code", control_code, 15),
            ("title", title, 24),
            ("description", description, 14),
        ]:
            value = _doc_lower(field_value)
            if value and value in lowered:
                score += weight
                reasons.append(f"Coincidencia directa con {field_name}.")

        title_terms = [
            term for term in _doc_lower(title).replace("/", " ").replace("-", " ").split()
            if len(term) >= 5
        ]
        matching_terms = [term for term in title_terms if term in lowered]

        if matching_terms:
            score += min(35, len(matching_terms) * 8)
            reasons.append(f"Términos relevantes detectados: {', '.join(matching_terms[:5])}.")

        if score >= 18:
            scored.append({
                "control_ref": control_code or clause or None,
                "standard_code": standard,
                "clause": clause or None,
                "title": title or None,
                "reason": " ".join(reasons) or "Coincidencia semántica preliminar con el contenido.",
                "confidence": round(min(0.94, 0.45 + score / 100), 2),
            })

    scored.sort(key=lambda item: item.get("confidence", 0), reverse=True)
    return scored[:8]


def _assess_missing_elements(document_type: str, text: str, metadata: Dict[str, Any]) -> List[str]:
    lowered = _doc_lower(text)
    missing = []

    if len(lowered) < 500:
        missing.append("El texto extraído es limitado; la revisión debe considerarse preliminar.")

    if not _doc_contains_any(lowered, ["versión", "version", "v.", "rev.", "revisión", "revision"]):
        missing.append("No se detecta versión o revisión documental explícita.")

    if not _doc_contains_any(lowered, ["aprobado", "aprobación", "aprobacion", "autoriza", "firma", "revisado por"]):
        missing.append("No se detecta aprobación formal o evidencia de revisión.")

    if not _doc_contains_any(lowered, ["responsable", "owner", "propietario", "encargado"]):
        missing.append("No se detecta responsable documental o dueño del proceso.")

    if not _doc_contains_any(lowered, ["fecha", "vigencia", "actualización", "actualizacion", "emisión", "emision"]):
        missing.append("No se detecta fecha de emisión, vigencia o actualización dentro del contenido.")

    if document_type in ["policy", "procedure"] and not _doc_contains_any(lowered, ["objetivo", "alcance", "scope"]):
        missing.append("No se detecta objetivo y/o alcance formal del documento.")

    if metadata.get("extraction", {}).get("truncated"):
        missing.append("El texto fue truncado por límite de análisis; revisar documento completo antes de aprobar.")

    return missing[:10]


def _detect_probable_errors(file_name: str, text: str, metadata: Dict[str, Any], standards: List[str]) -> List[str]:
    lowered = _doc_lower(text)
    errors = []

    modified_at = _doc_norm(metadata.get("modified_at"))
    if modified_at and modified_at[:4].isdigit():
        try:
            year = int(modified_at[:4])
            if year < 2023:
                errors.append("La fecha de modificación es antigua; validar si el documento sigue vigente.")
        except Exception:
            pass

    if "iso9001" in _doc_lower(file_name) and "ISO9001:2015" not in standards:
        errors.append("El nombre del archivo sugiere ISO9001, pero el análisis no logró confirmar esa norma con claridad.")

    if _doc_contains_any(lowered, ["borrador", "draft", "pendiente de aprobación", "pendiente de aprobacion"]):
        errors.append("El contenido sugiere estado borrador o pendiente de aprobación.")

    if _doc_contains_any(lowered, ["obsoleto", "no vigente", "reemplazado"]):
        errors.append("El contenido podría indicar documento obsoleto o reemplazado.")

    if _doc_contains_any(lowered, ["copia no controlada", "uncontrolled copy"]):
        errors.append("El contenido indica copia no controlada; no debe usarse como evidencia principal sin validación.")

    return errors[:10]


def _evidence_quality(text: str, document_type: str, missing_elements: List[str], probable_errors: List[str]) -> str:
    length = len(_doc_norm(text))

    if document_type == "unknown" or length < 300:
        return "insufficient"
    if probable_errors:
        return "medium"
    if len(missing_elements) >= 4:
        return "low"
    if len(missing_elements) >= 2:
        return "medium"
    if length >= 2500:
        return "high"

    return "medium"


def _confidence_score(
    text: str,
    standards: List[str],
    document_type: str,
    controls: List[Dict[str, Any]],
    missing_elements: List[str],
) -> float:
    score = 0.35

    if len(_doc_norm(text)) > 500:
        score += 0.14
    if len(_doc_norm(text)) > 2000:
        score += 0.10
    if standards:
        score += 0.14
    if document_type != "unknown":
        score += 0.12
    if controls:
        score += 0.08
    if len(missing_elements) <= 2:
        score += 0.07

    return round(max(0.1, min(0.95, score)), 2)


@app.post("/api/ai-compliance/analyze-document")
def analyze_document_for_compliance(
    payload: DocumentAnalysisRequest,
    x_ai_token: Optional[str] = Header(default=None),
    x_internal_token: Optional[str] = Header(default=None),
):
    """
    Analiza contenido documental extraído desde el Centro Inteligente de Evidencias.

    Reglas:
    - No aprueba cumplimiento.
    - No crea evidencia.
    - No modifica controles.
    - Devuelve recomendaciones revisables por humano.
    """
    configured = get_configured_ai_token()
    if configured:
        require_ai_token(
            x_ai_token=x_ai_token,
            x_internal_token=x_internal_token,
        )

    text = _doc_norm(payload.text)
    file_name = _doc_norm(payload.file_name)
    mime_type = _doc_norm(payload.mime_type)
    metadata = payload.metadata or {}

    standards = _infer_standards(file_name, text, payload.active_standards or [])
    document_type = _infer_document_type(file_name, mime_type, text)
    suggested_controls = _find_probable_controls(text, standards, payload.available_controls or [])
    missing_elements = _assess_missing_elements(document_type, text, metadata)
    probable_errors = _detect_probable_errors(file_name, text, metadata, standards)

    evidence_quality = _evidence_quality(text, document_type, missing_elements, probable_errors)
    confidence = _confidence_score(text, standards, document_type, suggested_controls, missing_elements)

    suggested_targets = [
        {
            "target_type": "control",
            "target_id": None,
            "standard_code": item.get("standard_code"),
            "control_ref": item.get("control_ref") or item.get("clause"),
            "reason": item.get("reason"),
            "confidence": item.get("confidence"),
        }
        for item in suggested_controls
    ]

    summary_parts = [
        f"Documento analizado como {document_type}.",
        f"Normas probables: {', '.join(standards) if standards else 'no determinada'}.",
        f"Calidad preliminar de evidencia: {evidence_quality}.",
    ]

    if probable_errors:
        summary_parts.append(f"Se detectaron {len(probable_errors)} advertencias que requieren revisión.")

    if suggested_controls:
        summary_parts.append(f"Se sugieren {len(suggested_controls)} posibles asociaciones a controles para revisión humana.")

    recommended_actions = [
        "Revisar el análisis antes de asociar el documento a controles o evidencias formales.",
        "Confirmar vigencia, responsable, versión y aprobación documental.",
        "Validar que el documento corresponda al alcance del tenant y a la norma activa aplicable.",
        "No considerar este resultado como aprobación de cumplimiento ni certificación.",
    ]

    if probable_errors:
        recommended_actions.append("Resolver advertencias detectadas antes de usar el documento como evidencia principal.")

    if suggested_controls:
        recommended_actions.append("Evaluar las asociaciones sugeridas y aprobar/rechazar manualmente cada una.")

    return {
        "document_type": document_type,
        "standards": standards,
        "suggested_controls": suggested_controls,
        "suggested_targets": suggested_targets,
        "summary": " ".join(summary_parts),
        "evidence_quality": evidence_quality,
        "missing_elements": missing_elements + probable_errors,
        "recommended_actions": recommended_actions,
        "confidence_score": confidence,
        "requires_review": True,
        "analysis_source": "ai_engine_document_auditor",
        "audit_findings": {
            "probable_errors": probable_errors,
            "missing_elements": missing_elements,
            "content_length": len(text),
            "metadata_reviewed": {
                "provider": metadata.get("provider"),
                "source_name": metadata.get("source_name"),
                "modified_at": metadata.get("modified_at"),
                "extraction_stage": metadata.get("extraction_stage"),
            },
        },
    }


@app.post("/semantic-evidence/analyze")
def analyze_semantic_evidence(
    payload: SemanticEvidenceAnalyzeRequest,
    x_ai_token: Optional[str] = Header(default=None),
    x_internal_token: Optional[str] = Header(default=None),
):
    """
    Sprint 3.5 semantic evidence endpoint.

    Rules:
    - Does not certify compliance.
    - Does not approve evidence.
    - Returns citeable fragments and reviewable suggestions only.
    """
    configured = get_configured_ai_token()
    if configured:
        require_ai_token(
            x_ai_token=x_ai_token,
            x_internal_token=x_internal_token,
        )

    text = _doc_norm(payload.text)
    filename = _doc_norm(payload.filename or payload.title)
    metadata = payload.metadata or {}
    document_type = _infer_document_type(filename, _doc_norm(metadata.get("mime_type")), text)
    standards = _infer_standards(filename, text, safe_list(metadata.get("active_standards")))
    missing_elements = _assess_missing_elements(document_type, text, metadata)
    probable_errors = _detect_probable_errors(filename, text, metadata, standards)
    evidence_quality = _evidence_quality(text, document_type, missing_elements, probable_errors)
    confidence = _confidence_score(text, standards, document_type, [], missing_elements)

    chunks: List[Dict[str, Any]] = []
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    for index in range(0, min(len(cleaned), 1200 * 12), 1200):
        chunk_text = cleaned[index:index + 1200].strip()
        if chunk_text:
            chunks.append({
                "chunk_index": len(chunks),
                "chunk_text": chunk_text,
                "page_number": None,
                "section_label": None,
                "hash": str(abs(hash(chunk_text))),
            })

    suggestions: List[Dict[str, Any]] = []
    haystack = _doc_lower(text)
    for group_name, targets in (payload.candidate_targets or {}).items():
        for target in targets[:80]:
            label = _doc_norm(target.get("label") or target.get("title") or target.get("name"))
            words = [word for word in _doc_lower(label).replace("-", " ").split() if len(word) >= 5]
            matches = [word for word in words if word in haystack]
            if not matches:
                continue
            score = min(0.94, 0.42 + len(matches) * 0.08)
            suggestions.append({
                "target_type": target.get("target_type") or group_name.rstrip("s"),
                "target_id": target.get("id") or target.get("target_id"),
                "target_label": label,
                "score": round(score, 2),
                "confidence": round(score, 2),
                "reason": f"Coincidencias detectadas: {', '.join(matches[:5])}. Requiere revision humana.",
                "chunk_index": 0 if chunks else None,
                "snippet": chunks[0]["chunk_text"][:500] if chunks else "",
            })

    return {
        "status": "processed" if text else "text_not_extractable",
        "classification": {
            "type": document_type,
            "confidence": confidence,
            "method": "rule_based",
            "reason": f"Calidad preliminar: {evidence_quality}. Revision humana obligatoria.",
        },
        "chunks": chunks,
        "suggestions": suggestions[:30],
        "scoring": {
            "relevance_to_object": round(confidence * 100, 2),
            "document_quality": evidence_quality,
            "traceability": 80 if chunks else 30,
            "human_review_required": True,
        },
    }
